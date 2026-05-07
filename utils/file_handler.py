"""
File I/O helpers for reading resumes and writing generated documents.
"""

from __future__ import annotations

import os
import re
from pathlib import Path

from utils.logger import get_logger

log = get_logger(__name__)


class FileHandler:
    """Read source resumes and persist generated DOCX/text files."""

    # ------------------------------------------------------------------ reads
    @staticmethod
    def read_pdf(path: str) -> str:
        """Extract all text from a PDF resume."""
        try:
            import PyPDF2  # noqa: PLC0415

            with open(path, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages)
            log.info("PDF loaded: %s (%d chars)", path, len(text))
            return text
        except Exception as exc:  # noqa: BLE001
            log.error("Failed to read PDF %s: %s", path, exc)
            raise

    @staticmethod
    def read_docx(path: str) -> str:
        """Extract all text from a DOCX resume."""
        try:
            from docx import Document  # noqa: PLC0415

            doc = Document(path)
            text = "\n".join(para.text for para in doc.paragraphs)
            log.info("DOCX loaded: %s (%d chars)", path, len(text))
            return text
        except Exception as exc:  # noqa: BLE001
            log.error("Failed to read DOCX %s: %s", path, exc)
            raise

    @staticmethod
    def read_resume(path: str) -> str:
        """Auto-detect format and return resume text."""
        lower = path.lower()
        if lower.endswith(".pdf"):
            return FileHandler.read_pdf(path)
        if lower.endswith(".docx"):
            return FileHandler.read_docx(path)
        # plain text fallback
        with open(path, encoding="utf-8") as fh:
            return fh.read()

    # ----------------------------------------------------------------- writes
    @staticmethod
    def save_text(content: str, path: str) -> None:
        """Write plain text to a file, creating parent directories as needed."""
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf-8") as fh:
            fh.write(content)
        log.info("Saved text → %s", path)

    @staticmethod
    def save_docx(content: str, path: str) -> None:
        """
        Convert plain-text / markdown content to a minimal DOCX and save it.
        Each line becomes a paragraph; lines starting with '#' become headings.
        """
        try:
            from docx import Document  # noqa: PLC0415
            from docx.shared import Pt  # noqa: PLC0415
        except ImportError as exc:
            raise ImportError("python-docx is required. Run: pip install python-docx") from exc

        Path(path).parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        for line in content.splitlines():
            stripped = line.strip()
            if stripped.startswith("###"):
                doc.add_heading(stripped.lstrip("#").strip(), level=3)
            elif stripped.startswith("##"):
                doc.add_heading(stripped.lstrip("#").strip(), level=2)
            elif stripped.startswith("#"):
                doc.add_heading(stripped.lstrip("#").strip(), level=1)
            else:
                para = doc.add_paragraph(stripped)
                para.style.font.size = Pt(11)
        doc.save(path)
        log.info("Saved DOCX → %s", path)

    @staticmethod
    def sanitise_filename(name: str) -> str:
        """Strip characters unsafe for file-system paths."""
        return re.sub(r"[^\w\-_. ]", "_", name).strip()
